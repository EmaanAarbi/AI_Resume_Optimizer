import streamlit as st
import google.generativeai as genai
import docx
import io
import os
from dotenv import load_dotenv

# --- Load environment variables ---
load_dotenv()

# --- Page config (must be first Streamlit call) ---
st.set_page_config(
    page_title="AI Resume Optimizer | Rewaa",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# --- Rewaa Design System CSS ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap');

/* ── Reset & base ── */
*, *::before, *::after { box-sizing: border-box; }

html, body, [data-testid="stAppViewContainer"] {
    font-family: 'DM Sans', sans-serif;
    background: #f4f6fa;
    color: #1a2332;
}

/* Hide default streamlit header */
[data-testid="stHeader"] { display: none; }
[data-testid="stToolbar"] { display: none; }
.block-container { padding-top: 0 !important; max-width: 760px !important; }

/* ── Top nav bar ── */
.rewaa-nav {
    background: #0b1f3a;
    padding: 14px 32px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin: -1rem -1rem 0 -1rem;
    width: calc(100% + 2rem);
}
.rewaa-nav .logo {
    color: #fff;
    font-size: 1.25rem;
    font-weight: 700;
    letter-spacing: -0.02em;
    display: flex;
    align-items: center;
    gap: 8px;
}
.rewaa-nav .logo-dot { color: #00d4aa; }
.rewaa-nav .badge {
    background: rgba(0,212,170,0.15);
    border: 1px solid rgba(0,212,170,0.3);
    color: #00d4aa;
    font-size: 0.7rem;
    font-weight: 600;
    padding: 3px 10px;
    border-radius: 20px;
    letter-spacing: 0.05em;
    text-transform: uppercase;
}

/* ── Hero section ── */
.hero {
    text-align: center;
    padding: 40px 24px 28px;
}
.hero h1 {
    font-size: 2rem;
    font-weight: 700;
    color: #0b1f3a;
    margin: 0 0 8px;
    letter-spacing: -0.03em;
    line-height: 1.2;
}
.hero p {
    color: #5a6a7e;
    font-size: 1rem;
    margin: 0;
    font-weight: 400;
}
.hero .accent { color: #00a884; }

/* ── Cards ── */
.card {
    background: #fff;
    border: 1px solid #e4e9f0;
    border-radius: 12px;
    padding: 24px;
    margin-bottom: 16px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}
.card-title {
    font-size: 0.75rem;
    font-weight: 700;
    color: #8a98a8;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    margin-bottom: 12px;
    display: flex;
    align-items: center;
    gap: 6px;
}
.card-title .step {
    background: #0b1f3a;
    color: #fff;
    width: 18px; height: 18px;
    border-radius: 50%;
    font-size: 0.65rem;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    font-weight: 700;
}

/* ── Status strip ── */
.status-strip {
    background: linear-gradient(135deg, #f0fdf9 0%, #e8f8f5 100%);
    border: 1px solid #c3e8e0;
    border-radius: 8px;
    padding: 12px 16px;
    display: flex;
    align-items: center;
    gap: 10px;
    margin-bottom: 16px;
    font-size: 0.875rem;
    color: #0a7c5c;
    font-weight: 500;
}
.status-strip .dot {
    width: 8px; height: 8px;
    border-radius: 50%;
    background: #00a884;
    flex-shrink: 0;
    animation: pulse 2s infinite;
}
@keyframes pulse {
    0%, 100% { opacity: 1; }
    50% { opacity: 0.4; }
}

/* ── Warning strip ── */
.warn-strip {
    background: #fffbeb;
    border: 1px solid #f5d879;
    border-radius: 8px;
    padding: 12px 16px;
    font-size: 0.875rem;
    color: #92610a;
    font-weight: 500;
    margin-bottom: 16px;
}

/* ── Result box ── */
.result-box {
    background: #f8fafc;
    border: 1px solid #e4e9f0;
    border-left: 4px solid #00a884;
    border-radius: 8px;
    padding: 20px;
    font-family: 'DM Mono', monospace;
    font-size: 0.82rem;
    line-height: 1.7;
    color: #2d3748;
    white-space: pre-wrap;
    word-break: break-word;
    max-height: 500px;
    overflow-y: auto;
}

/* ── Streamlit widget overrides ── */
[data-testid="stFileUploader"] > div {
    border: 2px dashed #c8d5e3 !important;
    border-radius: 10px !important;
    background: #f8fafc !important;
    transition: border-color 0.2s;
}
[data-testid="stFileUploader"] > div:hover {
    border-color: #00a884 !important;
}

div.stTextArea > div > textarea {
    border: 1.5px solid #dde4ef !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.9rem !important;
    background: #fafbfd !important;
    transition: border-color 0.2s;
}
div.stTextArea > div > textarea:focus {
    border-color: #00a884 !important;
    box-shadow: 0 0 0 3px rgba(0,168,132,0.1) !important;
}

div.stTextInput > div > div > input {
    border: 1.5px solid #dde4ef !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    background: #fafbfd !important;
}
div.stTextInput > div > div > input:focus {
    border-color: #00a884 !important;
    box-shadow: 0 0 0 3px rgba(0,168,132,0.1) !important;
}

div.stButton > button {
    background: #0b1f3a !important;
    color: #fff !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
    padding: 10px 28px !important;
    transition: background 0.2s, transform 0.1s !important;
    width: 100%;
}
div.stButton > button:hover {
    background: #1a3558 !important;
    transform: translateY(-1px) !important;
}
div.stButton > button:active { transform: translateY(0) !important; }

div.stDownloadButton > button {
    background: #fff !important;
    color: #0b1f3a !important;
    border: 1.5px solid #c8d5e3 !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
    transition: all 0.2s !important;
}
div.stDownloadButton > button:hover {
    border-color: #00a884 !important;
    color: #00a884 !important;
    background: #f0fdf9 !important;
}

div.stSelectbox > div > div {
    border: 1.5px solid #dde4ef !important;
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    background: #fafbfd !important;
}

div.stCheckbox > label { font-size: 0.9rem; }

/* ── Divider ── */
.section-divider {
    border: none;
    border-top: 1px solid #edf0f5;
    margin: 20px 0;
}

/* ── Footer ── */
.rewaa-footer {
    text-align: center;
    padding: 24px 0 8px;
    color: #9aabb8;
    font-size: 0.78rem;
}
.rewaa-footer a { color: #00a884; text-decoration: none; }

/* ── Metric pills ── */
.metrics-row {
    display: flex;
    gap: 12px;
    margin-bottom: 16px;
}
.metric-pill {
    flex: 1;
    background: #fff;
    border: 1px solid #e4e9f0;
    border-radius: 10px;
    padding: 14px 16px;
    text-align: center;
}
.metric-pill .metric-val {
    font-size: 1.4rem;
    font-weight: 700;
    color: #0b1f3a;
    display: block;
}
.metric-pill .metric-label {
    font-size: 0.72rem;
    color: #8a98a8;
    text-transform: uppercase;
    letter-spacing: 0.05em;
    font-weight: 600;
}

/* ── Spinner override ── */
[data-testid="stSpinner"] { color: #00a884 !important; }

/* ── Expander ── */
[data-testid="stExpander"] {
    border: 1px solid #e4e9f0 !important;
    border-radius: 10px !important;
}
</style>
""", unsafe_allow_html=True)


# ── Helpers ──────────────────────────────────────────────────────────────────

def read_docx(file) -> str:
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


def read_pdf(file) -> str:
    try:
        import fitz  # PyMuPDF
        pdf_document = fitz.open(stream=file.read(), filetype="pdf")
        return "".join(page.get_text() for page in pdf_document)
    except ImportError:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(file)
            return "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        except Exception as e:
            st.error(f"PDF reading failed: {e}")
            return ""


def create_docx(text: str) -> io.BytesIO:
    buffer = io.BytesIO()
    document = docx.Document()

    # Style the document
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = docx.shared.Pt(11)

    for line in text.split("\n"):
        document.add_paragraph(line)

    document.save(buffer)
    buffer.seek(0)
    return buffer


def create_pdf(text: str) -> io.BytesIO:
    from fpdf import FPDF
    buffer = io.BytesIO()
    pdf = FPDF()
    pdf.add_page()
    font_path = "DejaVuSans.ttf"
    if os.path.exists(font_path):
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=11)
    else:
        pdf.set_font("Helvetica", size=11)
    pdf.set_margins(20, 20, 20)
    pdf.multi_cell(0, 6, text)
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


def get_api_key() -> str | None:
    """Priority: session state → env var"""
    if st.session_state.get("gemini_key"):
        return st.session_state["gemini_key"]
    return os.getenv("GEMINI_API_KEY") or None


def word_count(text: str) -> int:
    return len(text.split())


def configure_gemini(api_key: str):
    genai.configure(api_key=api_key)


MODELS = [
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
    "gemini-2.5-flash",
]


def call_gemini(prompt: str, api_key: str) -> str:
    configure_gemini(api_key)
    last_err = None
    for model_name in MODELS:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            if "429" in str(e) or "quota" in str(e).lower():
                last_err = e
                continue
            raise e
    raise RuntimeError(
        f"All Gemini models quota exceeded. Last error: {last_err}"
    )


# ── Session state defaults ────────────────────────────────────────────────────
for key, val in {
    "result_text": None,
    "resume_text": None,
    "gemini_key": "",
    "mode": "optimize",
}.items():
    if key not in st.session_state:
        st.session_state[key] = val


# ── UI ────────────────────────────────────────────────────────────────────────

# Nav bar
st.markdown("""
<div class="rewaa-nav">
  <div class="logo">
    <span>rewaa</span><span class="logo-dot">.</span>
  </div>
  <div class="badge">AI Tools</div>
</div>
""", unsafe_allow_html=True)

# Hero
st.markdown("""
<div class="hero">
  <h1>AI Resume <span class="accent">Optimizer</span></h1>
  <p>Upload your resume and get an AI-powered, job-tailored version in seconds.</p>
</div>
""", unsafe_allow_html=True)

# ── API key config ────────────────────────────────────────────────────────────
api_key = get_api_key()
if api_key:
    st.markdown('<div class="status-strip"><div class="dot"></div>Gemini API connected — ready to optimize</div>', unsafe_allow_html=True)
else:
    st.markdown('<div class="warn-strip">⚠️ No Gemini API key found. Set <code>GEMINI_API_KEY</code> in your <code>.env</code> file or enter it below.</div>', unsafe_allow_html=True)
    with st.expander("🔑 Enter Gemini API Key"):
        typed_key = st.text_input(
            "Gemini API Key",
            type="password",
            placeholder="AIza...",
            label_visibility="collapsed",
        )
        if typed_key:
            st.session_state["gemini_key"] = typed_key
            api_key = typed_key

# ── Step 1: Upload ────────────────────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step">1</span>Upload Your Resume</div>', unsafe_allow_html=True)
uploaded_file = st.file_uploader(
    "Drag & drop or browse",
    type=["pdf", "docx"],
    label_visibility="collapsed",
)
st.markdown('</div>', unsafe_allow_html=True)

# ── Step 2: Job Description ───────────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step">2</span>Job Description <span style="color:#b0bec5;font-weight:400;text-transform:none;font-size:0.7rem;">(optional — leave blank to review & reformat only)</span></div>', unsafe_allow_html=True)
job_description = st.text_area(
    "Paste the job description here",
    placeholder="Paste the full job description here to tailor your resume to this specific role…",
    height=180,
    label_visibility="collapsed",
)
st.markdown('</div>', unsafe_allow_html=True)

# ── Step 3: Options ───────────────────────────────────────────────────────────
st.markdown('<div class="card"><div class="card-title"><span class="step">3</span>Optimization Options</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    tone = st.selectbox(
        "Tone",
        ["Professional", "Confident", "Technical", "Creative"],
        help="Adjusts writing style",
    )
with col2:
    output_format = st.selectbox(
        "Focus area",
        ["Full resume", "Summary only", "Skills section", "Work experience"],
    )

ats_mode = st.checkbox(
    "ATS-optimized output (keyword-rich, scannable)",
    value=True,
)
st.markdown('</div>', unsafe_allow_html=True)

# ── Run ───────────────────────────────────────────────────────────────────────
can_run = uploaded_file is not None and api_key

if not uploaded_file:
    st.info("👆 Upload a resume to get started.")

btn_label = "✨ Optimize Resume" if job_description.strip() else "✨ Review & Reformat Resume"
if st.button(btn_label, disabled=not can_run):
    if not api_key:
        st.error("Please provide a Gemini API key.")
        st.stop()

    with st.spinner("Reading your resume…"):
        try:
            if uploaded_file.type == "application/pdf":
                resume_text = read_pdf(uploaded_file)
            else:
                resume_text = read_docx(uploaded_file)

            if not resume_text.strip():
                st.error("Could not extract text from the uploaded file. Please check the file and try again.")
                st.stop()

            st.session_state["resume_text"] = resume_text
        except Exception as e:
            st.error(f"Error reading file: {e}")
            st.stop()

    with st.spinner("AI is analyzing and optimizing your resume…"):
        ats_note = (
            "\n- Use ATS-friendly formatting: clear section headers, bullet points, keyword-rich language."
            if ats_mode else ""
        )

        if job_description.strip():
            prompt = f"""You are an expert resume coach and career strategist.
Your task: rewrite and optimize the resume below to match the job description provided.

Guidelines:
- Tone: {tone}
- Focus: {output_format}
- Highlight matching skills, keywords, and achievements
- Quantify results wherever possible (use placeholder numbers if needed, marked with [X])
- Remove irrelevant experience{ats_note}
- Keep the output clean and professional

Job Description:
{job_description}

Original Resume:
{resume_text}

Provide the optimized resume only — no preamble, no commentary."""
        else:
            prompt = f"""You are an expert resume writer.
Review and reformat the resume below to make it more professional and impactful.

Guidelines:
- Tone: {tone}
- Focus: {output_format}
- Improve clarity, structure, and impact
- Fix grammar and phrasing issues
- Ensure consistent formatting{ats_note}

Original Resume:
{resume_text}

Return only the improved resume — no commentary."""

        try:
            result = call_gemini(prompt, api_key)
            st.session_state["result_text"] = result
        except Exception as e:
            st.error(f"AI error: {e}")
            st.stop()

# ── Output ────────────────────────────────────────────────────────────────────
if st.session_state["result_text"]:
    result_text = st.session_state["result_text"]
    original_text = st.session_state.get("resume_text", "")

    st.markdown("<hr class='section-divider'>", unsafe_allow_html=True)
    st.markdown("### ✅ Optimized Resume")

    # Stats row
    orig_words = word_count(original_text) if original_text else 0
    new_words = word_count(result_text)
    delta = new_words - orig_words
    delta_str = f"+{delta}" if delta >= 0 else str(delta)

    st.markdown(f"""
    <div class="metrics-row">
      <div class="metric-pill">
        <span class="metric-val">{orig_words}</span>
        <span class="metric-label">Original words</span>
      </div>
      <div class="metric-pill">
        <span class="metric-val">{new_words}</span>
        <span class="metric-label">Optimized words</span>
      </div>
      <div class="metric-pill">
        <span class="metric-val" style="color:#00a884">{delta_str}</span>
        <span class="metric-label">Word delta</span>
      </div>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("📄 View Full Optimized Resume", expanded=True):
        st.markdown(f'<div class="result-box">{result_text}</div>', unsafe_allow_html=True)

    # Edit & copy
    with st.expander("✏️ Edit before downloading"):
        edited = st.text_area(
            "Edit the resume text",
            value=result_text,
            height=400,
            label_visibility="collapsed",
        )
        if edited != result_text:
            if st.button("Save edits"):
                st.session_state["result_text"] = edited
                st.rerun()
        result_text = edited  # use edited version for downloads

    st.markdown("**Download as:**")
    col1, col2 = st.columns(2)
    with col1:
        docx_buf = create_docx(result_text)
        st.download_button(
            "📝 DOCX (Word)",
            data=docx_buf,
            file_name="optimized_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with col2:
        try:
            pdf_buf = create_pdf(result_text)
            st.download_button(
                "📄 PDF",
                data=pdf_buf,
                file_name="optimized_resume.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            st.warning(f"PDF export unavailable: {e}")

    if st.button("🔄 Start over / Optimize another"):
        st.session_state["result_text"] = None
        st.session_state["resume_text"] = None
        st.rerun()

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="rewaa-footer">
  Powered by <strong>Rewaa AI Tools</strong> &nbsp;·&nbsp;
  <a href="https://www.rewaatech.com" target="_blank">rewaatech.com</a>
  &nbsp;·&nbsp; © 2025 Rewaa Technology
</div>
""", unsafe_allow_html=True)
